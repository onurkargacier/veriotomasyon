import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.figure_factory as ff
import io
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches
import os
import seaborn as sns
import matplotlib.pyplot as plt

# Scikit-learn kütüphanelerini import edelim
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.compose import ColumnTransformer
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix, r2_score, mean_absolute_error

# Güvenlik ayarları
st.set_page_config(
    page_title="🤖 Veri Analiz Asistanı",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Basit erişim kontrolü (opsiyonel)
def check_access():
    """Basit erişim kontrolü - production'da daha güçlü bir sistem kullanın"""
    allowed_ips = os.getenv('ALLOWED_IPS', '').split(',')
    if allowed_ips and allowed_ips[0]:  # Eğer IP kısıtlaması varsa
        client_ip = st.experimental_get_query_params().get('client_ip', [''])[0]
        if client_ip not in allowed_ips:
            st.error("Erişim reddedildi. Lütfen sistem yöneticinizle iletişime geçin.")
            st.stop()

# Erişim kontrolünü çağır
check_access()

# --- Gelişmiş Tasarım için CSS (İlham Alınan Tasarım) ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');

    /* General Styling */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    .main {
        background-color: #ffffff;
    }
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }

    /* Hero Section */
    .hero {
        padding: 4rem 1rem;
        background-color: #f8f9fa;
        border-radius: 16px;
        text-align: center;
        margin-bottom: 2rem;
    }
    .hero h1 {
        font-size: 3.5rem;
        font-weight: 700;
        color: #212529;
        margin-bottom: 1rem;
    }
    .hero p {
        font-size: 1.25rem;
        color: #6c757d;
        margin-bottom: 2rem;
    }
    .hero .stFileUploader {
        max-width: 600px;
        margin: 0 auto;
        border: 2px dashed #6a0dad;
        background-color: #f3e8ff;
    }

    /* Feature Cards */
    .features {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
        gap: 1.5rem;
        text-align: left;
    }
    .feature-card {
        background-color: #f8f9fa;
        padding: 2rem;
        border-radius: 12px;
        border: 1px solid #e9ecef;
    }
    .feature-card h3 {
        color: #6a0dad; /* Purple */
        font-size: 1.5rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
     .feature-card p {
        color: #495057;
        font-size: 1rem;
    }

    /* Analysis View Styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
        justify-content: center;
    }
    .stTabs [data-baseweb="tab"] {
        font-weight: 600;
        background-color: #e9ecef;
        border-radius: 8px 8px 0 0;
        padding: 0.75rem 2rem;
        transition: all 0.3s;
    }
    .stTabs [data-baseweb="tab"]:hover {
        background-color: #dee2e6;
        color: #6a0dad;
    }
    .stTabs [aria-selected="true"] {
        background-color: #6a0dad;
        color: white;
    }
    .stButton>button {
        background-color: #fd7e14; /* Orange */
        color: white;
        border-radius: 8px;
        border: none;
        padding: 0.75rem 1.5rem;
        font-weight: 600;
    }
    .st-expander header {
        font-size: 1.3rem;
        font-weight: 700;
        color: #6a0dad;
    }
    .stMarkdown h2 {
        color: #6a0dad; 
        border-bottom: 2px solid #e9ecef;
        padding-bottom: 5px;
        margin-top: 2rem;
    }
     .stMarkdown h3 {
        color: #212529;
        margin-top: 1.5rem;
    }

    </style>
""", unsafe_allow_html=True)

def auto_categorize_dataset(df):
    text_cols = df.select_dtypes(include='object').columns
    num_cols = df.select_dtypes(include='number').columns
    date_cols = [col for col in df.columns if 'date' in col.lower() or 'time' in col.lower() or pd.api.types.is_datetime64_any_dtype(df[col])]
    summary = []
    if len(date_cols) > 0:
        summary.append("Zaman serisi analizi yapılabilir.")
    if len(num_cols) > 1:
        summary.append("Sayısal analizler, korelasyon ve regresyon yapılabilir.")
    if len(text_cols) > 0:
        summary.append("Kategorik veya metin analizi yapılabilir.")
    if len(num_cols) > 0 and len(text_cols) > 0:
        summary.append("Gruplar arası karşılaştırmalar yapılabilir.")
    if not summary:
        summary.append("Veri tipleri otomatik analiz için uygun değil.")
    return "\\n".join(summary)

# --- Session State ile Veri Yönetimi ---
if 'df' not in st.session_state:
    st.session_state.df = None
    st.session_state.uploaded_file_name = ""

# --- ANA UYGULAMA AKIŞI ---

if st.session_state.df is None:
    # Karşılama Ekranı
    st.markdown("""
    <div class="hero">
        <h1>Veri Analiz Asistanınız</h1>
        <p>Verilerinizi saniyeler içinde anlamlı içgörülere ve profesyonel raporlara dönüştürün.<br>
        Başlamak için dosyanızı yükleyin.</p>
    </div>
    """, unsafe_allow_html=True)
    
    up_file = st.file_uploader(
        "Analiz için CSV veya Excel dosyanızı buraya sürükleyin",
        type=["csv", "xlsx"],
        key="main_uploader"
    )

    if up_file:
        try:
            file_ext = up_file.name.split(".")[-1]
            df = pd.read_csv(up_file) if file_ext == "csv" else pd.read_excel(up_file)
            st.session_state.df = df
            st.session_state.uploaded_file_name = up_file.name
            st.rerun()
        except Exception as e:
            st.error(f"Dosya okunurken bir hata oluştu: {e}")

    st.markdown('<div style="margin-top: 3rem;"></div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="features">
        <div class="feature-card">
            <h3>✨ Otomatik Analiz</h3>
            <p>Veri setinizi otomatik olarak tarar, istatistiksel özetler ve analiz önerileri sunar.</p>
        </div>
        <div class="feature-card">
            <h3>Etkileşimli Grafikler</h3>
            <p>Verilerinizi görselleştirmek için dinamik ve modern grafikler (Histogram, Boxplot, Scatter vb.) oluşturur.</p>
        </div>
        <div class="feature-card">
            <h3>🤖 AI Destekli Modelleme</h3>
            <p>PyCaret entegrasyonu ile en iyi makine öğrenmesi modelini bulur ve tahminler yapar.</p>
        </div>
        <div class="feature-card">
            <h3>Anında Raporlama</h3>
            <p>Tüm analiz sonuçlarını tek bir tıkla indirilebilir bir Word raporuna dönüştürür.</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

else:
    # Analiz Arayüzü
    df = st.session_state.df
    uploaded_file_name = st.session_state.uploaded_file_name
    
    st.success(f"✅ **{uploaded_file_name}** üzerinde çalışılıyor.")
    
    if st.button("⬅️ Yeni Dosya Yükle"):
        st.session_state.df = None
        st.rerun()

    with st.expander("Veri Önizleme ve Sütun Tipleri", expanded=False):
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown("### Veri Önizlemesi")
            st.dataframe(df.head(), use_container_width=True)
        with col2:
            st.markdown("### Sütun Tipleri")
            st.dataframe(pd.DataFrame(df.dtypes, columns=["Tip"]).T, use_container_width=True)

    st.markdown("<h2>💡 Otomatik Analiz Önerisi</h2>", unsafe_allow_html=True)
    st.info(auto_categorize_dataset(df))

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "Temel Analizler",
        "📈 Gruplama Analizi",
        "🧮 Kategorik Analiz",
        "🤖 Otomatik Modelleme",
        "📄 Raporlama"
    ])

    # GRAFİKLERİ GEÇİCİ DOSYA YERİNE HAFIZADA TUTMAK İÇİN YARDIMCI FONKSİYON
    # Bu yöntem, diske veri yazmayı önleyerek güvenliği artırır.
    def fig_to_image_stream(fig):
        """Converts a Plotly figure to a PNG image in an in-memory byte stream."""
        buffer = io.BytesIO()
        fig.write_image(buffer, format="png", width=900, height=500, scale=2)
        buffer.seek(0)
        return buffer

    with tab1:
        st.markdown("<h2>📊 Temel İstatistikler ve Grafikler</h2>", unsafe_allow_html=True)
        with st.expander("Tanımlayıcı İstatistikler Tablosu"):
            st.dataframe(df.describe().T, use_container_width=True)

        num_cols = df.select_dtypes(include='number').columns.tolist()
        # Değişken adları, artık dosya yolu değil, hafızadaki stream'i tuttuğunu yansıtacak şekilde güncellendi.
        histogram_streams, boxplot_streams, scatter_stream, corr_stream = [], [], None, None
        
        if len(num_cols) > 0:
            st.markdown("<h3>Sayısal Değişken Analizi</h3>", unsafe_allow_html=True)
            selected_num = st.multiselect(
                "Analiz için sayısal sütun(lar) seçin:",
                num_cols,
                default=num_cols[:min(1, len(num_cols))],
                key="num_graphs"
            )

            if selected_num:
                for col in selected_num:
                    st.markdown(f"#### `{col}` Değişkeni Analizi")
                    c1, c2 = st.columns(2)
                    with c1:
                        fig1 = px.histogram(df, x=col, nbins=30, title=f"{col} Histogramı", color_discrete_sequence=["#6a0dad"])
                        fig1.update_layout(margin=dict(l=20, r=20, t=40, b=20), plot_bgcolor="white", title_x=0.5)
                        st.plotly_chart(fig1, use_container_width=True)
                        histogram_streams.append((col, fig_to_image_stream(fig1)))
                    with c2:
                        fig2 = px.box(df, y=col, title=f"{col} Boxplot", color_discrete_sequence=["#fd7e14"])
                        fig2.update_layout(margin=dict(l=20, r=20, t=40, b=20), plot_bgcolor="white", title_x=0.5)
                        st.plotly_chart(fig2, use_container_width=True)
                        boxplot_streams.append((col, fig_to_image_stream(fig2)))

            if len(selected_num) > 1:
                st.markdown("<h3>İlişki Analizi</h3>", unsafe_allow_html=True)
                c1, c2 = st.columns(2)
                with c1:
                    st.markdown("#### Dağılım Grafiği (Scatter Plot)")
                    x_col = st.selectbox("X ekseni:", selected_num, key="scatter_x")
                    y_col = st.selectbox("Y ekseni:", selected_num, index=min(1, len(selected_num)-1), key="scatter_y")
                    if x_col != y_col:
                        fig_scatter = px.scatter(df, x=x_col, y=y_col, title=f"{x_col} vs {y_col}", color_discrete_sequence=["#10b981"])
                        fig_scatter.update_layout(margin=dict(l=20, r=20, t=40, b=20), plot_bgcolor="white", title_x=0.5)
                        st.plotly_chart(fig_scatter, use_container_width=True)
                        scatter_stream = fig_to_image_stream(fig_scatter)
                    else:
                        st.warning("Farklı değişkenler seçerek dağılım grafiğini görebilirsiniz.")
                with c2:
                    st.markdown("#### Korelasyon Matrisi")
                    corr = df[selected_num].corr(numeric_only=True)
                    fig_corr = ff.create_annotated_heatmap(
                        z=corr.values, x=corr.columns.tolist(), y=corr.index.tolist(),
                        annotation_text=corr.round(2).values, colorscale="Purples"
                    )
                    fig_corr.update_layout(title="Korelasyon Matrisi", margin=dict(l=20, r=20, t=40, b=20), plot_bgcolor="white", title_x=0.5)
                    st.plotly_chart(fig_corr, use_container_width=True)
                    corr_stream = fig_to_image_stream(fig_corr)
        else:
             st.warning("Temel analizler için en az bir sayısal sütun gereklidir.")

    with tab2:
        st.markdown("<h2>📈 Gruplar Arası Karşılaştırma</h2>", unsafe_allow_html=True)
        cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        num_cols = df.select_dtypes(include='number').columns.tolist()
        group_boxplot_stream = None
        if cat_cols and num_cols:
            col1, col2 = st.columns(2)
            with col1:
                group_col = st.selectbox("Gruplama (kategorik):", cat_cols, key="group_col")
            with col2:
                value_col = st.selectbox("Değer (sayısal):", num_cols, key="value_col")
            
            fig = px.box(df, x=group_col, y=value_col, title=f"{group_col} Gruplarına Göre {value_col} Dağılımı", color=group_col, color_discrete_sequence=px.colors.qualitative.Pastel)
            fig.update_layout(margin=dict(l=20, r=20, t=40, b=20), plot_bgcolor="white", title_x=0.5)
            st.plotly_chart(fig, use_container_width=True)
            group_boxplot_stream = fig_to_image_stream(fig)
        else:
            st.warning("Gruplar arası karşılaştırma için en az bir kategorik ve bir sayısal sütun gereklidir.")

    with tab3:
        st.markdown("<h2>🧮 Kategorik Değişken Analizi</h2>", unsafe_allow_html=True)
        cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        barplot_streams = []
        if cat_cols:
            selected_cat = st.multiselect("Grafik için kategorik sütun(lar) seçin:", cat_cols, default=cat_cols[:min(1, len(cat_cols))], key="cat_graphs")
            for col in selected_cat:
                vc = df[col].value_counts().head(15)
                fig = px.bar(x=vc.index, y=vc.values, title=f"{col} - Frekans Dağılımı (En Sık 15)", color=vc.index, color_discrete_sequence=px.colors.qualitative.Set3)
                fig.update_layout(margin=dict(l=20, r=20, t=40, b=20), plot_bgcolor="white", title_x=0.5, xaxis_title=col, yaxis_title="Frekans")
                st.plotly_chart(fig, use_container_width=True)
                barplot_streams.append((col, fig_to_image_stream(fig)))
        else:
            st.warning("Kategorik analiz için uygun sütun bulunamadı.")

    with tab4:
        st.markdown("<h2>🤖 Otomatik Modelleme ve Tahmin</h2>", unsafe_allow_html=True)
        st.info("🎯 Hedef değişkeni seçin, sistem otomatik olarak en iyi modeli bulsun ve tahmin yapsın.")

        target = st.selectbox("Hedef (tahmin edilecek) değişkeni seçin:", df.columns.tolist(), index=len(df.columns)-1, key="target_select")

        df_model = df.dropna(subset=[target])
        problem_type = "classification"
        if pd.api.types.is_numeric_dtype(df_model[target]):
            problem_type = "regression"
            st.success("Problem Tipi: **Regresyon** (Sayısal bir değeri tahmin etme)")
        else:
            st.success("Problem Tipi: **Sınıflandırma** (Bir kategoriyi tahmin etme)")

        if st.button("🚀 Otomatik Modelleme Başlat", key="auto_ml_btn"):
            if not target:
                st.warning("Lütfen modelleme için bir hedef değişkeni seçin.")
            else:
                with st.spinner("Model için veriler hazırlanıyor ve eğitiliyor, lütfen bekleyin..."):
                    try:
                        # Adım 1: Veriyi Hazırlama
                        df_model = df.dropna(subset=[target])
                        X = df_model.drop(target, axis=1)
                        y = df_model[target]

                        # Sütunları sayısal ve kategorik olarak ayır
                        numeric_features = X.select_dtypes(include=['number']).columns
                        categorical_features = X.select_dtypes(include=['object', 'category']).columns

                        # Adım 2: Veri Ön İşleme Hattı (Pipeline) Oluşturma
                        # Bu hat, sayısal verileri ölçeklendirir ve kategorik verileri one-hot encoding yapar.
                        preprocessor = ColumnTransformer(
                            transformers=[
                                ('num', StandardScaler(), numeric_features),
                                ('cat', OneHotEncoder(handle_unknown='ignore'), categorical_features)
                            ])

                        # Adım 3: Model Seçimi ve Pipeline'ı Tamamlama
                        if problem_type == "regression":
                            model = RandomForestRegressor(n_estimators=100, random_state=42)
                            st.session_state.model_type = "Regresyon"
                        else: # classification
                            # Sınıflandırma için target'ın en az 2 sınıfı olduğundan emin ol
                            if y.nunique() < 2:
                                 st.error("Sınıflandırma için hedef değişkende en az 2 farklı sınıf olmalıdır.")
                                 st.stop()
                            model = RandomForestClassifier(n_estimators=100, random_state=42)
                            st.session_state.model_type = "Sınıflandırma"

                        # Ön işlemci ve modeli birleştir
                        pipeline = Pipeline(steps=[('preprocessor', preprocessor),
                                                  ('model', model)])

                        # Adım 4: Veriyi Eğitim ve Test Olarak Ayırma ve Modeli Eğitme
                        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
                        pipeline.fit(X_train, y_train)
                        
                        # Tahminleri yap
                        y_pred = pipeline.predict(X_test)
                        
                        st.session_state.model_pipeline = pipeline
                        st.session_state.X_test = X_test
                        st.session_state.y_test = y_test
                        st.session_state.y_pred = y_pred

                        st.success("✅ Modelleme başarıyla tamamlandı! Sonuçlar aşağıdadır.")

                    except Exception as e:
                        st.error(f"Modelleme sırasında bir hata oluştu: {e}")
                        st.stop()

    # --- Model Sonuçlarını Gösterme ---
    if 'model_pipeline' in st.session_state:
        st.markdown("<h3>🤖 Model Sonuçları ve Yorumlama</h3>", unsafe_allow_html=True)
        
        # Sonuçları getirelim
        pipeline = st.session_state.model_pipeline
        X_test = st.session_state.X_test
        y_test = st.session_state.y_test
        y_pred = st.session_state.y_pred
        model_type = st.session_state.model_type

        if model_type == "Regresyon":
            st.write("#### Regresyon Model Performansı")
            r2 = r2_score(y_test, y_pred)
            mae = mean_absolute_error(y_test, y_pred)
            st.metric("R-kare Skoru (R²)", f"{r2:.2f}")
            st.metric("Ortalama Mutlak Hata (MAE)", f"{mae:.2f}")
            
            st.write("R-kare, modelin hedef değişkendeki varyansın ne kadarını açıkladığını gösterir. 1'e ne kadar yakınsa o kadar iyidir. MAE, tahminlerin gerçek değerlerden ortalama ne kadar saptığını gösterir.")

            # Tahmin vs Gerçek Değer Grafiği
            fig_pred = px.scatter(x=y_test, y=y_pred, labels={'x': 'Gerçek Değerler', 'y': 'Tahmin Edilen Değerler'},
                                  title="Gerçek ve Tahmin Edilen Değerlerin Karşılaştırılması")
            fig_pred.add_shape(type='line', x0=y_test.min(), y0=y_test.min(), x1=y_test.max(), y1=y_test.max(), line=dict(color='Red'))
            st.plotly_chart(fig_pred, use_container_width=True)

        else: # Sınıflandırma
            st.write("#### Sınıflandırma Model Performansı")
            accuracy = accuracy_score(y_test, y_pred)
            st.metric("Doğruluk (Accuracy)", f"{accuracy:.2%}")
            
            st.text("Sınıflandırma Raporu:")
            report = classification_report(y_test, y_pred, output_dict=True)
            st.dataframe(pd.DataFrame(report).transpose())

            # Karışıklık Matrisi (Confusion Matrix)
            st.write("#### Karışıklık Matrisi (Confusion Matrix)")
            st.info("Bu matris, modelin hangi sınıfları doğru, hangilerini yanlış tahmin ettiğini gösterir. Köşegen üzerindeki sayılar doğru tahminlerdir.")
            cm = confusion_matrix(y_test, y_pred, labels=pipeline.classes_)
            fig_cm = plt.figure(figsize=(8, 6))
            sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=pipeline.classes_, yticklabels=pipeline.classes_)
            plt.xlabel('Tahmin Edilen Sınıf')
            plt.ylabel('Gerçek Sınıf')
            st.pyplot(fig_cm)

        # Özellik Önem Sıralaması (Feature Importance) - Her iki model için de ortak
        st.write("<h3>🎯 Sonucu En Çok Ne Etkiliyor? (Özellik Önem Sıralaması)</h3>", unsafe_allow_html=True)
        try:
            # Pipeline'dan one-hot encoder ile oluşturulan yeni sütun isimlerini alalım
            ohe_feature_names = pipeline.named_steps['preprocessor'].named_transformers_['cat'].get_feature_names_out(X.select_dtypes(include=['object', 'category']).columns)
            # Sayısal sütun isimlerini de ekleyelim
            all_feature_names = list(X.select_dtypes(include=['number']).columns) + list(ohe_feature_names)
            
            importances = pipeline.named_steps['model'].feature_importances_
            feature_importance_df = pd.DataFrame(list(zip(all_feature_names, importances)), columns=['Özellik', 'Önem']).sort_values('Önem', ascending=False)
            
            # En önemli 15 özelliği gösterelim
            fig_fi = px.bar(feature_importance_df.head(15), x='Önem', y='Özellik', orientation='h', title='En Önemli 15 Özellik')
            fig_fi.update_layout(yaxis={'categoryorder':'total ascending'})
            st.plotly_chart(fig_fi, use_container_width=True)
            st.info("Bu grafik, modelin tahmin yaparken hangi sütunlara daha fazla 'önem' verdiğini gösterir. Yüksek çubuklar, sonucun belirlenmesinde daha etkilidir.")
        except Exception as e:
            st.warning(f"Özellik önemi grafiği oluşturulamadı: {e}")

    with tab5:
        st.markdown("<h2>📄 Otomatik Raporlama ve İndir</h2>", unsafe_allow_html=True)
        st.info("Aşağıda analiz ve modelleme sonuçlarının özetini görebilir, raporu Word olarak indirebilirsiniz.")

        rapor_ozet = f"""
Otomatik Veri Analiz Raporu
Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}
Yüklenen dosya: {uploaded_file_name}
Satır sayısı: {df.shape[0]}
Sütun sayısı: {df.shape[1]}

Temel İstatistikler:
{df.describe().to_string()}

Otomatik Analiz Önerisi:
{auto_categorize_dataset(df)}
        """

        st.text_area("Rapor Özeti", rapor_ozet, height=300)

        if st.button("Word Raporu Oluştur ve İndir"):
            with st.spinner("Rapor oluşturuluyor, lütfen bekleyin..."):
                doc = Document()
                doc.add_heading('Otomatik Veri Analiz Raporu', 0)
                doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                doc.add_paragraph(f"Yüklenen dosya: {uploaded_file_name}")
                doc.add_paragraph(f"Satır sayısı: {df.shape[0]}, Sütun sayısı: {df.shape[1]}")
                doc.add_heading('Temel İstatistikler', level=1)
                doc.add_paragraph(df.describe().to_string())
                
                if histogram_streams:
                    doc.add_heading('Histogramlar', level=1)
                    for col, stream in histogram_streams:
                        doc.add_paragraph(f"{col} Histogramı", style='Heading 2')
                        doc.add_picture(stream, width=Inches(5.5))

                if boxplot_streams:
                    doc.add_heading('Boxplotlar', level=1)
                    for col, stream in boxplot_streams:
                        doc.add_paragraph(f"{col} Boxplot", style='Heading 2')
                        doc.add_picture(stream, width=Inches(5.5))

                if scatter_stream:
                    doc.add_heading('Scatter Plot', level=1)
                    doc.add_picture(scatter_stream, width=Inches(5.5))

                if corr_stream:
                    doc.add_heading('Korelasyon Matrisi', level=1)
                    doc.add_picture(corr_stream, width=Inches(5.5))

                if group_boxplot_stream:
                    doc.add_heading('Gruplar Arası Boxplot', level=1)
                    doc.add_picture(group_boxplot_stream, width=Inches(5.5))

                if barplot_streams:
                    doc.add_heading('Kategorik Barplotlar', level=1)
                    for col, stream in barplot_streams:
                        doc.add_paragraph(f"{col} Barplot", style='Heading 2')
                        doc.add_picture(stream, width=Inches(5.5))

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="✓ Word Raporunu İndir",
                    data=buffer,
                    file_name="otomatik_analiz_raporu.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )